import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="<3 ABNT", layout="wide", page_icon="🎓")

# --- FUNÇÕES TÉCNICAS ---
def configurar_margens(doc):
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)

def aplicar_formato_corpo(para, texto):
    para_format = para.paragraph_format
    para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format.line_spacing = 1.5
    para_format.first_line_indent = Cm(1.25)
    
    run = para.add_run(texto)
    run.font.name = 'Arial'
    run.font.size = Pt(12)

def aplicar_formato_citacao_longa(para, texto):
    para_format = para.paragraph_format
    para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_format.line_spacing = 1.0
    para_format.left_indent = Cm(4)
    para_format.space_before = Pt(12)
    para_format.space_after = Pt(12)
    
    run = para.add_run(texto)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

# --- DIVIDINDO A TELA EM DUAS COLUNAS ---
col_principal, col_apoio = st.columns([7, 3], gap="large")

# ==========================================
# LADO ESQUERDO: O APP DE FORMATAÇÃO
# ==========================================
with col_principal:
    st.title(" Formatador ABNT")
    st.write("Facilitando a vida do estudante, um parágrafo por vez.")

    tipo_texto = st.radio("O que você vai colar agora?", ["Texto Comum (Parágrafos)", "Citação Longa (Mais de 3 linhas)"])
    texto_input = st.text_area("Cole seu texto aqui:", height=200)

    if 'documento' not in st.session_state:
        st.session_state.documento = Document()
        configurar_margens(st.session_state.documento)
        st.session_state.historico = []

    # Botão de Adicionar
    if st.button("Adicionar ao Documento"):
        if texto_input.strip():
            p = st.session_state.documento.add_paragraph()
            if tipo_texto == "Texto Comum (Parágrafos)":
                aplicar_formato_corpo(p, texto_input)
            else:
                aplicar_formato_citacao_longa(p, texto_input)
            
            st.session_state.historico.append(tipo_texto)
            st.success("Adicionado com sucesso!")
        else:
            st.warning("Opa! Cole o texto antes de clicar.")

    st.divider()

    # Parte de Download
    if st.session_state.historico:
        st.write(f"📝 Itens já formatados: {len(st.session_state.historico)}")
        
        buffer = io.BytesIO()
        st.session_state.documento.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="📥 Baixar Trabalho Completo (.docx)",
            data=buffer,
            file_name="trabalho_formatado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        if st.button(" Limpar tudo e recomeçar"):
            st.session_state.documento = Document()
            configurar_margens(st.session_state.documento)
            st.session_state.historico = []
            st.rerun()

# ==========================================
# LADO DIREITO: TEXTO DE APOIO E PIX
# ==========================================
with col_apoio:
    st.header("🎓 Apoie o Projeto")
    
    # 1. QR CODE NO TOPO
    try:
        st.image("qrcode.png", caption="Escaneie para apoiar a Nico ou envie por Pix para 91983270175! ☕.")
    except:
        st.info("Espaço para o QR Code (qrcode.png)")
    
    # 2. TEXTO DE EXPLICAÇÃO
    st.markdown("""
    ### 🛠️ Apoie uma Engenheira em Formação!
    
    Sou a **Nico**, 25 anos, futura Engenheira de Controle e Automação pelo **IFPA** (9º semestre). Desenvolvi este formatador para devolver o tempo que a burocracia da ABNT rouba de nós.
    
    **Por que o seu apoio é imprescindível hoje?** Na Engenharia, a inovação não acontece sentada em uma mesa. Ela acontece no laboratório, na bancada de robótica e no campo. Atualmente, meu desenvolvimento está "preso" a um PC fixo, o que é um gargalo crítico na minha reta final de curso.
    
    Ter um notebook funcional não é um luxo, é a **condição básica** para eu levar meus códigos para o laboratório e entregar meu TCC. 
    
    Ao apoiar, você não está apenas fazendo uma doação; você está **investindo no futuro da tecnologia nacional** e ajudando uma estudante a cruzar a linha de chegada.
    """)
